from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Callable

from core.plugin_base import BasePlugin, PluginContext, PluginResult
from core.paths import UPLOADS_DIR
from core.safe_download import download_to_path


TEXT_DOCUMENT_SUFFIXES = {
    ".txt",
    ".md",
    ".markdown",
    ".json",
    ".py",
    ".sh",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".cfg",
    ".conf",
    ".csv",
    ".log",
}
IMAGE_DOCUMENT_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp"}
SAVE_ONLY_DOCUMENT_SUFFIXES = {".pdf", ".pptx", ".xlsx"}
DOCX_CONTEXT_LIMIT = 24000
WORD_FOLLOW_UP_INTENTS = (
    "总结",
    "分析",
    "提炼",
    "改写",
    "梳理",
    "生成ppt",
    "生成 ppt",
    "提取要点",
)


class DocumentIntakePlugin(BasePlugin):
    name = "document_intake"
    priority = 25

    def _raw(self, context: PluginContext) -> dict[str, Any]:
        raw = context.metadata.get("raw")
        return raw if isinstance(raw, dict) else {}

    def _load_content_obj(self, raw: dict[str, Any]) -> dict[str, Any]:
        content = raw.get("content")
        if isinstance(content, dict):
            return content
        if isinstance(content, str):
            try:
                loaded = json.loads(content)
            except Exception:
                return {"raw": content}
            return loaded if isinstance(loaded, dict) else {"raw": loaded}
        return {}

    def _find_values_by_key(self, obj: Any, wanted_keys: set[str]) -> list[str]:
        result: list[str] = []

        def walk(value: Any) -> None:
            if isinstance(value, dict):
                for key, item in value.items():
                    if str(key).lower() in wanted_keys and item:
                        result.append(str(item))
                    walk(item)
            elif isinstance(value, list):
                for item in value:
                    walk(item)

        walk(obj)
        return result

    def sanitize_filename(self, name: str | None, fallback: str = "dingtalk_file") -> str:
        clean = Path(str(name or fallback)).name.strip()
        clean = re.sub(r"[^A-Za-z0-9._() \-\u4e00-\u9fff]", "_", clean)
        clean = clean.strip(" .")
        return clean or fallback

    def extract_document_info(self, raw: dict[str, Any]) -> dict[str, Any]:
        content_obj = self._load_content_obj(raw)

        filename = ""
        for key in ("fileName", "filename", "name", "title"):
            value = content_obj.get(key) if isinstance(content_obj, dict) else None
            if isinstance(value, str) and value.strip():
                filename = value.strip()
                break

        if not filename:
            for key in ("fileName", "filename", "name", "title"):
                value = raw.get(key)
                if isinstance(value, str) and value.strip():
                    filename = value.strip()
                    break

        codes = self._find_values_by_key(content_obj, {"downloadcode", "download_code"})
        for key in ("downloadCode", "download_code"):
            if raw.get(key):
                codes.append(str(raw.get(key)))

        urls = self._find_values_by_key(content_obj, {"downloadurl", "download_url", "url"})
        for key in ("downloadUrl", "download_url", "url"):
            if raw.get(key):
                urls.append(str(raw.get(key)))

        return {
            "filename": self.sanitize_filename(filename or "dingtalk_file"),
            "download_codes": self._dedup(codes),
            "download_urls": self._dedup([url for url in urls if url.startswith("http")]),
            "file_id": content_obj.get("fileId") if isinstance(content_obj, dict) else raw.get("fileId"),
            "space_id": content_obj.get("spaceId") if isinstance(content_obj, dict) else raw.get("spaceId"),
        }

    def _dedup(self, values: list[str]) -> list[str]:
        deduped: list[str] = []
        for value in values:
            if value and value not in deduped:
                deduped.append(value)
        return deduped

    def is_text_suffix(self, suffix: str) -> bool:
        return (suffix or "").lower() in TEXT_DOCUMENT_SUFFIXES

    def is_image_suffix(self, suffix: str) -> bool:
        return (suffix or "").lower() in IMAGE_DOCUMENT_SUFFIXES

    def is_save_only_suffix(self, suffix: str) -> bool:
        return (suffix or "").lower() in SAVE_ONLY_DOCUMENT_SUFFIXES

    def is_docx_suffix(self, suffix: str) -> bool:
        return (suffix or "").lower() == ".docx"

    def is_legacy_doc_suffix(self, suffix: str) -> bool:
        return (suffix or "").lower() == ".doc"

    def classify_suffix(self, suffix: str) -> str:
        if self.is_text_suffix(suffix):
            return "text"
        if self.is_image_suffix(suffix):
            return "image"
        if self.is_docx_suffix(suffix):
            return "docx"
        if self.is_legacy_doc_suffix(suffix):
            return "legacy_doc"
        if self.is_save_only_suffix(suffix):
            return "save_only"
        return "unknown"

    def should_use_word_context(self, text: str) -> bool:
        normalized = (text or "").strip().lower()
        return any(intent in normalized for intent in WORD_FOLLOW_UP_INTENTS)

    def _cell_text(self, value: str) -> str:
        return " ".join((value or "").split()).replace("|", "\\|")

    def extract_docx_text(self, path: str | Path, limit: int = DOCX_CONTEXT_LIMIT) -> str:
        """Extract visible paragraphs and tables from a .docx into compact Markdown."""
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("未安装 python-docx，无法解析 .docx 文件。") from exc

        document = Document(str(path))
        sections: list[str] = []
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        paragraphs = [paragraph for paragraph in paragraphs if paragraph]
        if paragraphs:
            sections.append("\n".join(paragraphs))

        for table_index, table in enumerate(document.tables, start=1):
            rows = [
                [self._cell_text(cell.text) for cell in row.cells]
                for row in table.rows
            ]
            rows = [row for row in rows if any(row)]
            if not rows:
                continue
            width = max(len(row) for row in rows)
            normalized_rows = [row + [""] * (width - len(row)) for row in rows]
            header = normalized_rows[0]
            markdown = [
                f"表格 {table_index}",
                "| " + " | ".join(header) + " |",
                "| " + " | ".join(["---"] * width) + " |",
            ]
            markdown.extend("| " + " | ".join(row) + " |" for row in normalized_rows[1:])
            sections.append("\n".join(markdown))

        extracted = "\n\n".join(sections).strip()
        if not extracted:
            extracted = "（未提取到可见正文或表格内容）"
        if len(extracted) > limit:
            extracted = extracted[:limit].rstrip() + f"\n\n[已截断，原始解析文本超过 {limit} 字符]"
        return extracted

    def _safe_suffix(self, content_type: str, fallback: str = ".bin") -> str:
        ctype = (content_type or "").lower()
        if "jpeg" in ctype or "jpg" in ctype:
            return ".jpg"
        if "png" in ctype:
            return ".png"
        if "webp" in ctype:
            return ".webp"
        if "json" in ctype:
            return ".json"
        if "markdown" in ctype:
            return ".md"
        if "text" in ctype:
            return ".txt"
        if "pdf" in ctype:
            return ".pdf"
        return fallback

    def _save_raw_message(self, raw: dict[str, Any], raw_upload_dir: Path) -> str:
        raw_upload_dir.mkdir(parents=True, exist_ok=True)
        raw_path = raw_upload_dir / f"raw_file_msg_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}.json"
        raw_path.write_text(json.dumps(raw, ensure_ascii=False, indent=2), encoding="utf-8")
        return str(raw_path)

    def _download_url_to_raw_file(self, url: str, filename: str, raw_upload_dir: Path) -> str:
        raw_upload_dir.mkdir(parents=True, exist_ok=True)
        safe_name = self.sanitize_filename(filename)
        if not Path(safe_name).suffix:
            safe_name = f"{safe_name}.bin"

        path = raw_upload_dir / f"{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}_{safe_name}"
        download_to_path(url, path)
        return str(path)

    def _download_code_to_raw_file(
        self,
        code: str,
        filename: str,
        raw_upload_dir: Path,
        download_url_resolver: Callable[[str], str | None] | None,
    ) -> str:
        if download_url_resolver is None:
            raise RuntimeError("未提供钉钉 downloadCode 下载 URL 解析器。")

        url = download_url_resolver(code)
        if not url:
            raise RuntimeError("downloadCode 解析结果为空，无法下载文件。")
        return self._download_url_to_raw_file(url, filename, raw_upload_dir)

    def match(self, context: PluginContext) -> bool:
        return self._raw(context).get("msgtype") == "file"

    def handle(self, context: PluginContext) -> PluginResult:
        raw = self._raw(context)
        if raw.get("msgtype") != "file":
            return PluginResult(
                handled=False,
                metadata={"plugin": self.name, "reason": "not_file_msgtype"},
            )

        raw_upload_dir = Path(
            context.metadata.get("raw_upload_dir")
            or str(UPLOADS_DIR / "dingtalk_raw")
        )
        download_url_resolver = context.metadata.get("download_url_resolver")

        raw_path = self._save_raw_message(raw, raw_upload_dir)
        info = self.extract_document_info(raw)
        filename = info.get("filename") or "dingtalk_file"
        errors: list[str] = []
        saved_path = ""

        for code in info.get("download_codes") or []:
            try:
                saved_path = self._download_code_to_raw_file(
                    code,
                    filename,
                    raw_upload_dir,
                    download_url_resolver,
                )
                break
            except Exception as exc:
                errors.append(str(exc))

        if not saved_path:
            for url in info.get("download_urls") or []:
                try:
                    saved_path = self._download_url_to_raw_file(url, filename, raw_upload_dir)
                    break
                except Exception as exc:
                    errors.append(str(exc))

        base_metadata = {
            "plugin": self.name,
            "mode": "document_intake",
            "raw_path": raw_path,
            "filename": filename,
            "download_codes": info.get("download_codes") or [],
            "download_urls": info.get("download_urls") or [],
        }

        if not saved_path:
            return PluginResult(
                handled=True,
                text="文件消息已收到，但下载失败。",
                metadata={
                    **base_metadata,
                    "document_kind": "download_failed",
                    "reason": "download_failed",
                    "errors": errors[-5:] or ["未拿到可用下载码或下载 URL。"],
                },
            )

        suffix = Path(saved_path).suffix.lower()
        document_kind = self.classify_suffix(suffix)
        metadata = {
            **base_metadata,
            "document_kind": document_kind,
            "saved_path": saved_path,
            "suffix": suffix,
        }

        if document_kind == "text":
            try:
                file_text = Path(saved_path).read_text(encoding="utf-8")
            except UnicodeDecodeError as exc:
                return PluginResult(
                    handled=True,
                    text=f"文件已保存：{saved_path}\n\n但不是有效 UTF-8 文本，暂未解析：{exc}",
                    metadata={
                        **metadata,
                        "document_kind": "text_decode_failed",
                        "reason": "text_decode_failed",
                    },
                )

            return PluginResult(
                handled=True,
                text=file_text,
                files=[saved_path],
                metadata={
                    **metadata,
                    "file_text": file_text,
                    "reason": "text_file",
                },
            )

        if document_kind == "image":
            return PluginResult(
                handled=True,
                files=[saved_path],
                metadata={**metadata, "reason": "image_file"},
            )

        if document_kind == "legacy_doc":
            return PluginResult(
                handled=True,
                text="已收到 Word 文件，但老式 .doc 暂不支持解析。请先另存为 .docx 后重新上传。",
                files=[saved_path],
                metadata={**metadata, "reason": "legacy_doc_unsupported"},
            )

        if document_kind == "docx":
            sidecar_path = str(Path(saved_path).with_suffix(".txt"))
            try:
                file_text = self.extract_docx_text(saved_path)
                Path(sidecar_path).write_text(file_text, encoding="utf-8")
            except Exception as exc:
                return PluginResult(
                    handled=True,
                    text=(
                        f"已收到 Word 文件并保留原文件：{saved_path}\n\n"
                        f"但文本解析失败：{exc}"
                    ),
                    files=[saved_path],
                    metadata={
                        **metadata,
                        "document_kind": "docx_parse_failed",
                        "reason": "docx_parse_failed",
                    },
                )

            return PluginResult(
                handled=True,
                text="已收到 Word 文件，并已解析出文本。你可以继续让我总结、改写、提炼要点或生成 PPT。",
                files=[saved_path, sidecar_path],
                metadata={
                    **metadata,
                    "file_text": file_text,
                    "sidecar_path": sidecar_path,
                    "reason": "docx_parsed",
                },
            )

        if document_kind == "save_only":
            return PluginResult(
                handled=True,
                text=f"文件已收到并保存：{saved_path}\n\n当前暂未解析该类型文件。",
                files=[saved_path],
                metadata={**metadata, "reason": "save_only"},
            )

        return PluginResult(
            handled=True,
            text=f"文件已收到并保存：{saved_path}\n\n当前暂未解析该类型文件。",
            files=[saved_path],
            metadata={**metadata, "reason": "unknown_suffix"},
        )
